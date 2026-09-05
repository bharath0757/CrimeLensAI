"""Preview contract tests: real parsing, explicit substitute at the NLP HTTP boundary."""

import hashlib
import json
from io import BytesIO

import httpx
import pytest
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from app.api.deps import get_ai_service
from app.core.config import settings
from app.integrations.ai_integration import IntegratedAIService
from app.main import app


@pytest.fixture
def preview_service():
    calls = []

    def handler(request):
        body = json.loads(request.content)
        calls.append(body)
        text = body["text"]
        mentions = []
        if "9123456789" in text:
            start = text.index("9123456789")
            mentions.append({
                "entity_id": "phone-mention", "entity_type": "PHONE", "value": "9123456789",
                "normalized_value": "+919123456789", "confidence": .98,
                "start_offset": start, "end_offset": start + 10,
                "source_field": "document_text", "case_id": body["case_id"],
            })
        return httpx.Response(200, json={"entities": mentions, "model": "test-contract", "warnings": []})

    service = IntegratedAIService(transport=httpx.MockTransport(handler))
    app.dependency_overrides[get_ai_service] = lambda: service
    yield calls
    app.dependency_overrides.pop(get_ai_service, None)


def test_preview_requires_authentication(client, preview_service):
    assert client.post("/api/v1/extraction/preview", json={"text": "9123456789"}).status_code == 401
    assert client.post("/api/v1/extraction/preview-file", files={"file": ("fir.txt", b"9123456789")}).status_code == 401
    assert preview_service == []


def test_preview_preserves_unicode_offsets_without_creating_case(client, admin_auth_headers, preview_service):
    count = client.get("/api/v1/cases", headers=admin_auth_headers).json()["total"]
    narrative = "Witness 📞 reported 9123456789."
    response = client.post("/api/v1/extraction/preview", json={"text": narrative}, headers=admin_auth_headers)
    assert response.status_code == 200
    result = response.json()
    mention = result["entities"][0]
    assert result["text"][mention["start_offset"]:mention["end_offset"]] == "9123456789"
    assert result["document_sha256"] is None
    assert preview_service[0]["case_id"] is None
    assert client.get("/api/v1/cases", headers=admin_auth_headers).json()["total"] == count


@pytest.mark.parametrize("text", ["", "   ", "x" * 500001])
def test_preview_rejects_empty_and_oversized_text(client, admin_auth_headers, preview_service, text):
    assert client.post("/api/v1/extraction/preview", json={"text": text}, headers=admin_auth_headers).status_code == 422
    assert preview_service == []


def test_uploaded_txt_is_decoded_and_hashed(client, admin_auth_headers, preview_service):
    raw = b"\xef\xbb\xbfWitness called 9123456789."
    result = client.post("/api/v1/extraction/preview-file", headers=admin_auth_headers, files={"file": ("unknown-fir.txt", raw)})
    assert result.status_code == 200
    assert result.json()["document_sha256"] == hashlib.sha256(raw).hexdigest()
    assert preview_service[0]["text"] == "Witness called 9123456789."


def test_real_docx_paragraphs_and_tables_are_read(client, admin_auth_headers, preview_service):
    document = Document()
    document.add_paragraph("FIR witness narrative")
    document.add_table(rows=1, cols=1).cell(0, 0).text = "Phone: 9123456789"
    stream = BytesIO()
    document.save(stream)
    response = client.post("/api/v1/extraction/preview-file", headers=admin_auth_headers, files={"file": ("fir.docx", stream.getvalue())})
    assert response.status_code == 200
    assert response.json()["text"] == "FIR witness narrative\nPhone: 9123456789"
    assert len(response.json()["entities"]) == 1


def test_real_text_pdf_is_read(client, admin_auth_headers, preview_service):
    writer = PdfWriter()
    page = writer.add_blank_page(width=595, height=842)
    font = DictionaryObject({NameObject("/Type"): NameObject("/Font"), NameObject("/Subtype"): NameObject("/Type1"), NameObject("/BaseFont"): NameObject("/Helvetica")})
    page[NameObject("/Resources")] = DictionaryObject({NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)})})
    content = DecodedStreamObject()
    content.set_data(b"BT /F1 12 Tf 40 800 Td (Phone: 9123456789) Tj ET")
    page[NameObject("/Contents")] = writer._add_object(content)
    stream = BytesIO()
    writer.write(stream)
    response = client.post("/api/v1/extraction/preview-file", headers=admin_auth_headers, files={"file": ("fir.pdf", stream.getvalue())})
    assert response.status_code == 200
    assert "9123456789" in preview_service[0]["text"]
    assert len(response.json()["entities"]) == 1


@pytest.mark.parametrize("filename,raw", [("bad.pdf", b"not a pdf"), ("scan.png", b"image"), ("empty.txt", b""), ("bad.txt", b"\xff")])
def test_unreadable_files_never_reach_nlp(client, admin_auth_headers, preview_service, filename, raw):
    response = client.post("/api/v1/extraction/preview-file", headers=admin_auth_headers, files={"file": (filename, raw)})
    assert response.status_code == 422
    assert preview_service == []


def test_oversized_upload_is_rejected(client, admin_auth_headers, preview_service, monkeypatch):
    monkeypatch.setattr(settings, "MAX_FILE_SIZE_BYTES", 10)
    response = client.post("/api/v1/extraction/preview-file", headers=admin_auth_headers, files={"file": ("big.txt", b"x" * 11)})
    assert response.status_code == 413
    assert preview_service == []


def test_upstream_failure_is_clear_and_does_not_leak_url(client, admin_auth_headers):
    def fail(_):
        raise httpx.ConnectError("private internal URL secret")
    service = IntegratedAIService(transport=httpx.MockTransport(fail))
    app.dependency_overrides[get_ai_service] = lambda: service
    try:
        response = client.post("/api/v1/extraction/preview", json={"text": "9123456789"}, headers=admin_auth_headers)
        assert response.status_code == 503
        assert "private internal" not in response.text
    finally:
        app.dependency_overrides.pop(get_ai_service, None)
