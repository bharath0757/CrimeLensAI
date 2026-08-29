import json
import sys
from pathlib import Path

from docx import Document


def extract(path: Path) -> dict:
    doc = Document(path)
    blocks = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append({
                "type": "paragraph",
                "style": paragraph.style.name if paragraph.style else None,
                "text": text,
            })
    tables = []
    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({"table": table_index, "rows": rows})
    return {
        "path": str(path),
        "paragraphs": blocks,
        "tables": tables,
        "inline_shapes": len(doc.inline_shapes),
        "sections": len(doc.sections),
    }


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    for filename in sys.argv[1:]:
        print(json.dumps(extract(Path(filename)), ensure_ascii=False, indent=2))
